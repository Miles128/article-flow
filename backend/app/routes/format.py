"""格式转换路由 - 使用配置外部化 + @with_llm 装饰器"""
import io
import re
from flask import Blueprint, request, jsonify, send_file
from ..decorators import with_llm
from ..config_loader import get_platforms
from ..utils import translate_error

bp = Blueprint('format', __name__)


@bp.route('/inline-css', methods=['POST'])
def inline_css():
    data = request.json
    if not data or 'html' not in data or 'css' not in data:
        return jsonify({'error': 'html and css are required'}), 400

    try:
        from premailer import transform
        full_html = f"""<!DOCTYPE html>
<html><head><style>{data['css']}</style></head>
<body>{data['html']}</body></html>"""

        inlined = transform(full_html, remove_classes=False, strip_important=True)

        body_match = re.search(r'<body>(.*?)</body>', inlined, re.DOTALL)
        result = body_match.group(1) if body_match else inlined

        return jsonify({'inlined_html': result})
    except Exception as e:
        return jsonify({'error': translate_error(e)}), 500


@bp.route('/normalize', methods=['POST'])
@with_llm(require_content=False)
def normalize_markdown(llm, data):
    if 'content' not in data:
        return jsonify({'error': 'content is required'}), 400

    result = llm.chat([
        {'role': 'system', 'content': '你是一位Markdown格式专家，擅长统一文档格式。'},
        {'role': 'user', 'content': f'请将以下Markdown内容规范化，确保格式统一：\n\n{data["content"]}\n\n规范化要求：\n1. 标题层级统一\n2. 列表格式统一\n3. 引用格式统一\n4. 代码块格式统一\n5. 空行统一\n\n请直接输出规范化后的Markdown内容。'}
    ])
    return {'normalized_content': result}


@bp.route('/convert', methods=['POST'])
@with_llm(require_content=False)
def convert_format(llm, data):
    if 'content' not in data or 'target_platform' not in data:
        return jsonify({'error': 'content and target_platform are required'}), 400

    converted = llm.convert_format(data['content'], data['target_platform'])
    return {'target_platform': data['target_platform'], 'converted_content': converted}


@bp.route('/platforms', methods=['GET'])
def get_platforms_api():
    platforms = get_platforms()
    return jsonify(platforms)


@bp.route('/export', methods=['POST'])
def export_content():
    data = request.json
    if not data or 'content' not in data or 'format' not in data:
        return jsonify({'error': 'content and format are required'}), 400

    content = data['content']
    export_format = data['format']
    title = data.get('title', 'Untitled').replace('/', '-')[:60]
    safe_filename = title if title else 'untitled'

    if export_format == 'html':
        html_content = markdown_to_html(content)
        return jsonify({
            'format': 'html',
            'content': html_content,
            'filename': f"{safe_filename}.html",
        })

    elif export_format == 'pdf':
        buf = markdown_to_pdf(content, title)
        return send_file(
            buf,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'{safe_filename}.pdf',
        )

    elif export_format == 'docx':
        buf = markdown_to_docx(content)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{safe_filename}.docx',
        )

    elif export_format == 'plaintext':
        plain = markdown_to_plaintext(content)
        return jsonify({
            'format': 'plaintext',
            'content': plain,
            'filename': f"{safe_filename}.txt",
        })

    return jsonify({'error': f'Unknown format: {export_format}'}), 400


def markdown_to_html(md_content):
    import markdown
    html = markdown.markdown(md_content, extensions=[
        'extra', 'codehilite', 'toc', 'tables', 'fenced_code'
    ])

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Exported Article</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; line-height: 1.8; max-width: 800px; margin: 0 auto; padding: 40px 20px; color: #333; }}
        h1, h2, h3 {{ margin-top: 1.5em; margin-bottom: 0.5em; font-weight: 600; }}
        h1 {{ font-size: 2em; border-bottom: 2px solid #eee; padding-bottom: 0.3em; }}
        h2 {{ font-size: 1.5em; }}
        h3 {{ font-size: 1.25em; }}
        p {{ margin: 1em 0; }}
        blockquote {{ border-left: 4px solid #ddd; margin: 1em 0; padding-left: 1em; color: #666; }}
        code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 4px; }}
        pre {{ background: #f4f4f4; padding: 1em; border-radius: 8px; overflow-x: auto; }}
        pre code {{ background: none; padding: 0; }}
        ul, ol {{ padding-left: 2em; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px 12px; }}
        th {{ background: #f4f4f4; }}
        img {{ max-width: 100%; height: auto; }}
        a {{ color: #0366d6; text-decoration: none; }}
        a:hover {{ text-decoration: underline; }}
    </style>
</head>
<body>
{html}
</body>
</html>"""


def markdown_to_plaintext(md_content):
    text = md_content
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'```[\s\S]*?```', '', text)
    text = re.sub(r'^\s*[-*+]\s+', '• ', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*\d+\.\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'^>\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ─── Markdown → 块级结构解析 ───

def _parse_md_blocks(md: str) -> list[dict]:
    """将 Markdown 文本解析为块级结构列表。"""
    blocks: list[dict] = []
    lines = md.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        # 标题
        m = re.match(r'^(#{1,6})\s+(.+)', line)
        if m:
            blocks.append({'kind': 'heading', 'level': len(m.group(1)), 'text': m.group(2).strip()})
            i += 1
            continue
        # 无序列表
        m = re.match(r'^(\s*)[-*+]\s+(.+)', line)
        if m:
            items = [m.group(2).strip()]
            indent = len(m.group(1))
            i += 1
            while i < len(lines) and re.match(rf'^\s{{{indent + 2},}}[-*+]\s+', lines[i]):
                items.append(re.sub(r'^\s+[-*+]\s+', '', lines[i]).strip())
                i += 1
            blocks.append({'kind': 'ul', 'items': items})
            continue
        # 有序列表
        m = re.match(r'^(\s*)\d+\.\s+(.+)', line)
        if m:
            items = [m.group(2).strip()]
            i += 1
            while i < len(lines) and re.match(r'^\s*\d+\.\s+', lines[i]):
                items.append(re.sub(r'^\s*\d+\.\s+', '', lines[i]).strip())
                i += 1
            blocks.append({'kind': 'ol', 'items': items})
            continue
        # 空行
        if not line.strip():
            i += 1
            continue
        # 普通段落
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r'^(#{1,6}\s|[-*+]\s|\d+\.\s)', lines[i]):
            para_lines.append(lines[i])
            i += 1
        blocks.append({'kind': 'para', 'text': '\n'.join(para_lines).strip()})
    return blocks


def _strip_inline(text: str) -> str:
    """去除行内 Markdown 标记，返回纯文本。"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'!\[.*?\]\([^)]+\)', '', text)
    return text


# ─── DOCX 导出 ───

def markdown_to_docx(md_content: str) -> io.BytesIO:
    """将 Markdown 内容转换为 DOCX 文件（python-docx）。"""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    blocks = _parse_md_blocks(md_content)
    for b in blocks:
        kind = b['kind']
        if kind == 'heading':
            level = min(b['level'], 3)
            p = doc.add_heading(_strip_inline(b['text']), level=level)
        elif kind == 'para':
            doc.add_paragraph(_strip_inline(b['text']))
        elif kind in ('ul', 'ol'):
            for item in b['items']:
                p = doc.add_paragraph(_strip_inline(item), style='List Bullet')
        else:
            doc.add_paragraph(_strip_inline(b.get('text', '')))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── PDF 导出 ───

def markdown_to_pdf(md_content: str, title: str = '') -> io.BytesIO:
    """将 Markdown 内容转换为 PDF 文件（fpdf2）。"""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # 注册中文字体（使用内置的 helvetica 作为后备；fpdf2 支持 Unicode）
    # fpdf2 内置支持 CJK 需要用 add_font 注册支持 Unicode 的字体
    # 使用 fpdf2 自带的 Unicode 字体机制
    try:
        pdf.add_font('NotoSansSC', '', '/System/Library/Fonts/Supplemental/Arial Unicode.ttf')
        font_name = 'NotoSansSC'
    except Exception:
        # Fallback: use built-in (no CJK, but won't crash)
        font_name = 'Helvetica'

    pdf.set_font(font_name, '', 12)

    if title:
        pdf.set_font(font_name, '', 16)
        pdf.multi_cell(0, 10, title, align='C')
        pdf.ln(6)

    blocks = _parse_md_blocks(md_content)
    pdf.set_font(font_name, '', 11)

    for b in blocks:
        kind = b['kind']
        if kind == 'heading':
            size = {1: 16, 2: 14, 3: 12}.get(b['level'], 11)
            pdf.set_font(font_name, '', size)
            pdf.ln(4)
            pdf.multi_cell(0, 8, _strip_inline(b['text']))
            pdf.set_font(font_name, '', 11)
            pdf.ln(2)
        elif kind == 'para':
            pdf.multi_cell(0, 6, _strip_inline(b['text']))
            pdf.ln(2)
        elif kind in ('ul', 'ol'):
            for idx, item in enumerate(b['items']):
                prefix = '• ' if kind == 'ul' else f'{idx + 1}. '
                pdf.multi_cell(0, 6, prefix + _strip_inline(item))
            pdf.ln(2)
        else:
            pdf.multi_cell(0, 6, _strip_inline(b.get('text', '')))
            pdf.ln(2)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf
